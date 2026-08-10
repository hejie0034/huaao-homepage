from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\huawei\Desktop\小蜜蜂Agent")
OUT = ROOT / "小蜜蜂agent交付报告.docx"

FONT_CN = "宋体"
FONT_HEAD = "黑体"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "E7E6E6"
BORDER_GRAY = "BFBFBF"


def set_run_font(run, name=FONT_CN, size=10.5, bold=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), BORDER_GRAY)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=GRAY)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=9.5)
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    hdr = table.rows[0]
    prevent_row_split(hdr)
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True, size=10)
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for i, value in enumerate(values):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT_CN
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
normal.font.size = Pt(10.5)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

heading_tokens = {
    "Heading 1": (16, 18, 10),
    "Heading 2": (13, 14, 7),
    "Heading 3": (12, 10, 5),
}
for name, (size, before, after) in heading_tokens.items():
    style = styles[name]
    style.font.name = FONT_HEAD
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = BLACK
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = FONT_CN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run("小蜜蜂agent交付报告  |  第 ")
set_run_font(r, size=9, color=GRAY)
add_page_field(footer_p)
r = footer_p.add_run(" 页")
set_run_font(r, size=9, color=GRAY)

# Title block: simple and deliberately undecorated.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)
run = title.add_run("小蜜蜂agent交付报告")
set_run_font(run, name=FONT_HEAD, size=20, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(3)
run = subtitle.add_run("项目说明、知识库说明与后续维护指南")
set_run_font(run, size=11, color=GRAY)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(14)
run = meta.add_run("版本：当前本地测试版    统计日期：2026年8月10日")
set_run_font(run, size=9.5, color=GRAY)

add_body(doc, "这份文档用于项目交接。接手者不需要先阅读代码，按照本文顺序即可了解小蜜蜂Agent的定位、已实现能力、知识来源、问题清单位置、启动方法、维护方法和当前不足。")

add_heading(doc, "先看这四个位置", 1)
add_numbers(doc, [
    r"要查看系统能回答的全部问题：打开 C:\Users\huawei\Desktop\小蜜蜂Agent\小蜜蜂全部问题清单.xlsx。",
    r"要启动系统：进入 C:\Users\huawei\Desktop\小蜜蜂Agent，运行 .codex_start_server.ps1。",
    r"要修改文字问答：查看 ulearning_teacher_faq.json 和 manual_2026_faq.json。",
    r"要维护图片或视频：查看“截图教程”“视频教程”和“指导视频”目录。",
])

add_heading(doc, "1. 项目基本信息", 1)
add_table(doc, ["项目项", "当前情况"], [
    ("项目名称", "小蜜蜂Agent"),
    ("服务对象", "uLearning平台教师"),
    ("项目目录", r"C:\Users\huawei\Desktop\小蜜蜂Agent"),
    ("默认地址", "http://127.0.0.1:8012/"),
    ("前端", "原生HTML、CSS、JavaScript"),
    ("后端", "Python ThreadingHTTPServer"),
    ("大模型", "DeepSeek，当前配置模型为 deepseek-chat"),
    ("当前阶段", "可演示、可试用的教师操作指导测试版，尚不是正式生产系统"),
], [2700, 6660])

add_heading(doc, "2. 项目定位", 1)
add_body(doc, "小蜜蜂Agent是一个面向uLearning教师的操作指导助手。它重点解决“入口在哪里、下一步点什么、为什么学生看不到、操作完成后怎样确认”等实际问题。")
add_body(doc, "系统的核心目标不是自由生成一篇长答案，而是给出教师能够照着完成的操作步骤，并在有资料时同时提供图片和视频。")
add_bullets(doc, [
    "主要用户：使用ULMS、Uclass及相关教师功能的老师、助教和课程管理员。",
    "主要任务：课程、班级、教学团队、课件资源、作业、讨论、课堂、考试、题库、成绩和AI工作台操作指导。",
    "回答方式：优先提供入口位置、编号步骤、完成检查和必要提醒。",
    "对话边界：允许正常闲聊和通用交流，但回答后会回到uLearning助手身份。",
])

add_heading(doc, "3. 已经实现的功能", 1)
add_heading(doc, "3.1 聊天和历史记录", 2)
add_bullets(doc, [
    "支持新建对话、输入问题、多行输入、发送状态和错误提示。",
    "支持用户问题重新编辑、复制；支持助手回答复制和重新生成。",
    "对话历史保存在当前浏览器的localStorage中，刷新页面后可以恢复。",
    "历史记录尚未接入账号系统，不能跨浏览器、跨电脑同步。",
])

add_heading(doc, "3.2 教程展示", 2)
add_bullets(doc, [
    "答案统一提供“文字解答、图片教程、视频教程”三个页签。",
    "图片教程左侧显示步骤，右侧显示截图，支持点击步骤和左右翻页。",
    "同一教程中的截图使用统一白色画布，切换时页面不会忽大忽小。",
    "截图页数严格等于真实截图数量，不再复制图片凑步骤。",
    "没有图片或视频时保留统一页签，并提示暂未提供。",
])

add_heading(doc, "3.3 智能问答和推荐", 2)
add_bullets(doc, [
    "系统区分闲聊、知识范围询问、平台操作问题、故障问题和未收录问题。",
    "平台操作问题只采用最强匹配答案，避免把多个相似流程混在一起。",
    "回答后推荐2至3个相关的下一步或相邻功能，尽量排除当前问题本身。",
    "未命中知识库的平台问题会提示用户补充页面、按钮和报错，并记录到 unanswered_questions.jsonl。",
])

add_heading(doc, "3.4 中英文", 2)
add_bullets(doc, [
    "左下角可一键切换中文和English，侧边栏宽度保持不变。",
    "页面按钮、输入提示、教程页签、反馈窗口和推荐问题支持英文。",
    "英文问题会先转换为中文检索，再把答案和推荐问题翻译成英文。",
    "历史记录不会随语言切换自动翻译；图片和视频暂时沿用中文版。",
])

add_heading(doc, "3.5 反馈", 2)
add_bullets(doc, [
    "支持点赞、点踩和转人工反馈。",
    "反馈保存在 feedback 目录的Excel表中。",
    "当前“转人工”是提交反馈记录，并没有真正连接在线客服坐席或工单系统。",
])

add_heading(doc, "4. 现有知识库", 1)
add_body(doc, "下面的统计来自当前运行中的项目，而不是旧版审计报告。")
add_table(doc, ["知识内容", "数量", "当前用途"], [
    ("基础教师FAQ", "237条", "正式参与文字问答检索"),
    ("2026教师手册FAQ", "51条", "正式参与文字问答检索；其中ULMS 46条、Uclass 5条"),
    ("可检索问答合计", "288条", "后端 /api/chat 的主要知识来源"),
    ("截图教程", "96组 / 320张", "在图片教程页签中展示"),
    ("视频教程", "29个", "在视频教程页签中播放"),
], [3000, 1700, 4660])

add_heading(doc, "4.1 已接入正式回答流程的知识", 2)
add_bullets(doc, [
    "ulearning_teacher_faq.json：基础教师问答。",
    "manual_2026_faq.json：从ULMS和Uclass 2026教师手册整理出的问答。",
    "manual_2026_guides.json：2026手册图片教程及中英文步骤。",
    "screenshot_guide_translations_en.json：原截图教程英文步骤。",
    "截图教程目录：前端实际展示的步骤图片。",
    "视频教程和指导视频目录：前端实际播放的视频。",
])

add_heading(doc, "4.2 已保存但尚未接入在线检索的资料", 2)
add_bullets(doc, [
    "help知识库.docx及其构建产物。",
    "切片读取知识库目录中的 knowledge_chunks.jsonl 和 knowledge_base.md。",
    "网页截图知识库、uLearning产品白皮书等源文档。",
    "不完全前置条件功能.txt；该文件目前为空。",
])
add_body(doc, "当前健康接口显示 helpKnowledgeCount、chunkKnowledgeCount、screenshotKnowledgeCount 和 preconditionKnowledgeCount 均为0。因此这些资料不能描述为已经在线检索使用。当前运行模式是 deepseek_faq_text_only。")

add_heading(doc, "5. 全部问题在哪里查看", 1)
add_body(doc, "人工查看全部问题时，以以下Excel为准：")
add_code_line(doc, r"C:\Users\huawei\Desktop\小蜜蜂Agent\小蜜蜂全部问题清单.xlsx")
add_table(doc, ["工作表", "问题数", "说明"], [
    ("图文视频齐全", "21", "同时有文字、图片和视频"),
    ("单项媒体教程", "104", "有文字，并配图片或视频中的一种"),
    ("仅文字解答", "163", "当前没有图片和视频"),
], [2800, 1500, 5060])
add_body(doc, "表格可按问题种类、教程形式、截图数、适用系统和来源进行筛选。它适合产品、运营和知识库维护人员检查覆盖范围。")
add_body(doc, "系统运行时实际读取的是以下JSON：")
add_bullets(doc, [
    r"C:\Users\huawei\Desktop\小蜜蜂Agent\ulearning_teacher_faq.json",
    r"C:\Users\huawei\Desktop\小蜜蜂Agent\manual_2026_faq.json",
])
add_body(doc, "简单理解：Excel用于人看和统计，JSON用于程序检索和回答。修改Excel不会自动修改运行中的JSON。")

add_heading(doc, "6. 系统怎样回答一个问题", 1)
add_numbers(doc, [
    "用户在网页输入问题。",
    "后端判断问题属于闲聊、知识范围询问、平台操作、故障排查，还是未收录问题。",
    "平台操作问题在288条FAQ中检索，只保留最强匹配结果。",
    "后端把结构化答案整理成入口位置、编号步骤、完成检查和注意事项。",
    "前端继续匹配对应的截图教程和视频教程。",
    "系统根据FAQ关系和工作流生成相关问题按钮。",
    "英文模式会把英文问题转换为中文检索，再翻译答案和推荐问题。",
    "平台相关问题没有命中时，系统要求补充信息并写入 unanswered_questions.jsonl。",
])

add_heading(doc, "7. 主动排障能力", 1)
add_body(doc, "系统已经具备基于关键词和规则的主动排障。教师说“学生看不到作业”时，系统不会直接重复布置作业步骤，而会先检查前置条件。")
add_bullets(doc, [
    "作业看不到：检查发布状态、班级、学生入班、时间和账号。",
    "考试或测验看不到：检查发布对象、考试时间、班级和状态。",
    "课件、章节或单元看不到：检查发布、关联课件、班级和可见状态。",
    "资源、视频或文档看不到：检查资源状态、位置、权限和发布时间。",
    "公告看不到：检查是否发布、发布班级和学生课程。",
    "按钮或入口缺失：检查账号角色、页面位置、负责班级和电脑端/App差异。",
    "课程或班课缺失：检查学校机构、加入班级、学生名单和学期。",
])
add_body(doc, "当前能力边界：它只会提出检查问题，不能直接读取uLearning后台的真实课程、权限、学生名单或发布状态，也不会记住每一个“是/否”回答形成完整诊断树。")

add_heading(doc, "8. 启动和使用", 1)
add_heading(doc, "8.1 启动", 2)
add_numbers(doc, [
    r"打开 C:\Users\huawei\Desktop\小蜜蜂Agent。",
    "在PowerShell中进入该目录。",
    "运行启动脚本或直接运行Python后端。",
    "浏览器打开 http://127.0.0.1:8012/。",
])
add_code_line(doc, r"powershell -ExecutionPolicy Bypass -File .\.codex_start_server.ps1")
add_body(doc, "也可以直接运行：")
add_code_line(doc, r"python .\web_agent.py --port 8012")

add_heading(doc, "8.2 检查服务", 2)
add_code_line(doc, r"http://127.0.0.1:8012/api/health")
add_body(doc, "正常时应返回 ok=true，并显示FAQ、截图和视频数量。当前预期为288条FAQ、96组截图教程、29个视频教程。")

add_heading(doc, "8.3 配置说明", 2)
add_bullets(doc, [
    "DeepSeek密钥保存在项目根目录的.env文件中。",
    "交付报告不写入真实密钥。",
    ".env已经列入.gitignore，不应上传到公开仓库。",
    "服务目前只监听127.0.0.1，只能在当前电脑访问。",
])

add_heading(doc, "9. 知识库维护方法", 1)
add_heading(doc, "9.1 新增或修改文字问答", 2)
add_numbers(doc, [
    "先检索是否已有相同或相近问题，避免创建两个答案版本。",
    "在对应FAQ JSON中新增或修改条目。",
    "填写id、category、question、summary、keywords、answer和related。",
    "answer至少包含entry、steps、check和risk。",
    "步骤使用清晰的1、2、3操作，不写只有一句话的笼统回答。",
    "related配置2至4个前置、下一步或常见分支，不能推荐当前问题本身。",
])

add_heading(doc, "9.2 新增图片教程", 2)
add_numbers(doc, [
    "在截图教程目录建立与功能对应的文件夹。",
    "按实际操作顺序放入step1、step2等图片。",
    "每张图片准备准确的步骤说明；没有人工说明时可先OCR，再人工校正。",
    "截图数量就是可翻页数量，不能复制同一张图片凑步骤。",
    "尺寸较小的图片应补白到统一画布，不能拉伸变形。",
    "同步检查中文问题、英文问题和教程别名是否能命中同一个教程。",
])

add_heading(doc, "9.3 新增视频教程", 2)
add_numbers(doc, [
    "把视频放入视频教程或指导视频目录。",
    "文件名使用清楚的功能名称。",
    "确认视频问题和FAQ问题使用相同意图，例如“快速创建课程”和“创建课程”应共用教程。",
    "运行媒体校验，确认文件路径有效。",
])

add_heading(doc, "9.4 重新导入2026教师手册", 2)
add_code_line(doc, r"python .\scripts\import_2026_teacher_manuals.py --ulms .\build_cache\manual_import\ulms.docx --uclass .\build_cache\manual_import\uclass.docx")
add_body(doc, "上面的命令使用当前缓存中的手册文件；如果手册移动了，应把两个路径替换为新文件的完整路径。导入后需要检查FAQ数量、图片页数、中文步骤、英文步骤和旧文件清理情况，不应只看脚本是否执行成功。")

add_heading(doc, "9.5 修改后必须运行的检查", 2)
add_code_line(doc, r"python .\scripts\validate_faq.py")
add_code_line(doc, r"python .\scripts\validate_tutorial_media.py")
add_code_line(doc, r"python .\scripts\test_chat_routing.py")
add_code_line(doc, r"python .\scripts\test_troubleshooting.py")

add_heading(doc, "10. 当前不足和已知风险", 1)
add_table(doc, ["问题", "当前情况", "影响"], [
    ("FAQ结构不完整", "133条缺少check和risk", "回答缺少完成确认和风险提醒"),
    ("步骤过短", "115条不足两个明确步骤", "部分回答仍不够适合第一次使用的老师"),
    ("媒体覆盖不足", "163个问题只有文字", "图文视频三合一覆盖率不足"),
    ("英文混排", "41个英文图片步骤仍含中文按钮名", "英文版阅读体验不统一"),
    ("储备知识未接入", "切片、help和前置条件知识计数均为0", "已有资料没有参与在线检索"),
    ("排障深度有限", "规则式检查，无真实状态读取和多轮诊断树", "不能自动确认故障根因"),
    ("数据架构", "无账号系统、数据库和正式客服", "不能支持正式多用户运营"),
    ("部署", "仅本机HTTP服务", "不能作为正式线上系统访问"),
    ("安全", "缺少鉴权、限流、请求大小限制和完整隐私机制", "不适合直接公网开放"),
    ("并发", "反馈写入Excel，未见明确写锁", "多人同时提交时可能冲突"),
    ("测试", "有脚本测试，但无完整CI和全量中英文回归", "修改后仍可能出现回归"),
    ("版本管理", "大量最新文件尚未形成完整Git提交", "GitHub不能视为本地最新完整版本"),
], [2300, 3600, 3460])

add_heading(doc, "11. 版本和交付状态", 1)
add_body(doc, "截至本报告统计时，主项目与Agent-TA中的核心程序、FAQ、手册问答和翻译文件内容一致。这个结论只表示本地两个目录的核心文件一致，不代表GitHub仓库已经包含全部本地成果。")
add_heading(doc, "11.1 交付时必须保留", 2)
add_bullets(doc, [
    "index.html、app.js、styles.css、web_agent.py。",
    "deepseek_prompt.md及智能体提示词文件。",
    "ulearning_teacher_faq.json、manual_2026_faq.json、manual_2026_guides.json。",
    "manual_2026_translations_en.json、screenshot_guide_translations_en.json。",
    "截图教程、视频教程、指导视频目录。",
    "scripts目录中的导入、校验、审计和测试脚本。",
    "小蜜蜂全部问题清单.xlsx。",
    "本报告。",
])

add_heading(doc, "11.2 后续接手检查清单", 2)
add_numbers(doc, [
    "确认项目可以从新电脑完整启动。",
    "确认 /api/health 返回288、96和29三个当前基准数量。",
    "确认全部问题清单Excel与两个FAQ JSON数量一致。",
    "补齐133条缺少check和risk的FAQ。",
    "修正115条步骤不足的问题。",
    "清理重复问题和冲突别名。",
    "重新生成覆盖288条FAQ和96组截图的审计报告。",
    "修复英文图片步骤中的中英文混排。",
    "把确认后的全部成果提交并推送到GitHub。",
    "正式上线前补充数据库、用户认证、安全、日志、监控和并发控制。",
])

add_heading(doc, "12. 当前测试结论", 1)
add_table(doc, ["检查项", "结果", "说明"], [
    ("服务健康检查", "通过", "模型已配置，FAQ和媒体均能加载"),
    ("教程媒体校验", "通过", "96组截图、320张图片、29个视频路径有效"),
    ("聊天分流测试", "通过", "闲聊、通用问答、排障和教程不串线"),
    ("主动排障测试", "通过", "8类故障场景和3类普通操作场景通过"),
    ("FAQ完整性校验", "未通过", "仍有133条缺少check/risk，115条步骤不足"),
], [3000, 1600, 4760])

add_heading(doc, "13. 最后说明", 1)
add_body(doc, "小蜜蜂Agent已经达到可演示、可试用、可继续扩充教师操作知识的阶段。它的主要产品形态、问答流程、图文视频展示、主动排障和中英文界面已经形成。下一阶段的重点不是继续增加零散界面功能，而是统一知识质量、补齐媒体、接入储备知识、完善测试和版本管理，并完成正式部署所需的用户、数据库、安全和监控能力。")
add_body(doc, "后续维护人员如果只记住一件事：任何新增功能都必须同时检查文字答案、图片步骤、视频匹配、英文说明、相关推荐和自然语言命中，不能只新增一个文件就认为功能已经完成。")

doc.core_properties.title = "小蜜蜂agent交付报告"
doc.core_properties.subject = "项目说明、知识库说明与后续维护指南"
doc.core_properties.author = "小蜜蜂Agent项目组"
doc.core_properties.keywords = "小蜜蜂Agent,uLearning,交付报告,知识库,维护指南"
doc.save(OUT)
print(OUT)
